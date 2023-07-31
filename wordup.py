import docx
from PyQt5.QtCore import QFile, QIODevice, Qt, QTextStream
from PyQt5.QtGui import QFont, QIcon, QTextCursor
from PyQt5.QtWidgets import (QAction, QDockWidget, QFileDialog, QFontDialog,
                             QLabel, QMainWindow, QMenu, QPlainTextEdit,
                             QPushButton, QStatusBar, QTextBrowser, QToolBar,
                             QVBoxLayout, QWidget, QHBoxLayout)
import sys
import openai
from PyQt5.QtWidgets import QApplication, QMainWindow, QTextEdit, QAction, QFileDialog, QFontDialog, QStatusBar, QMenu, QCheckBox, QLabel, QPushButton, QVBoxLayout, QHBoxLayout, QWidget, QPlainTextEdit, QTextBrowser
from PyQt5.QtCore import pyqtSignal
from PyQt5.QtWidgets import QStyleFactory
from PyQt5.QtCore import pyqtSlot
from bs4 import BeautifulSoup
import requests

from transformers import GPT2Tokenizer
from PyQt5.QtWidgets import (QApplication, QWidget, QVBoxLayout, QHBoxLayout, QTextEdit, QPushButton,
                             QLabel, QPlainTextEdit, QTabWidget, QDialog, QLineEdit, QCheckBox,
                             QTextBrowser, QMenu, QComboBox, QSpinBox, QMessageBox)
from PyQt5.QtGui import QTextDocument, QTextImageFormat
from PyQt5.QtCore import QSettings, pyqtSlot
from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import QFontDialog, QMessageBox
from spellchecker import SpellChecker
from PyQt5.QtGui import QTextCursor, QTextBlockFormat
from PyQt5.QtGui import QTextOption
import pyphen
from PyQt5.QtWidgets import QScrollBar
from PyQt5.QtGui import QTextCursor
from PyQt5.QtGui import QFontMetricsF
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QTextOption, QFontMetricsF, QPainter, QPixmap
from PyQt5.QtPrintSupport import QPrinter
from PyQt5.QtWidgets import QMainWindow, QDockWidget, QScrollArea
from PyQt5.QtCore import QUrl
from PyQt5.QtGui import QTextCursor, QDesktopServices, QImage, QPainter, QPen, QGuiApplication,  QTextCharFormat
from PyQt5.QtWidgets import QTextBrowser, QInputDialog, QLineEdit, QSizePolicy
from PyQt5.QtCore import QUrl
from io import BytesIO
import tempfile
import os
from PyQt5.QtWidgets import QPushButton, QVBoxLayout, QLabel, QDialog, QDialogButtonBox
from PyQt5.QtWidgets import QShortcut
from PyQt5.QtGui import QKeySequence, QTextCursor
from PyQt5.QtGui import QTextCharFormat
import os
from PyQt5.QtWidgets import QFileDialog, QMenuBar
from docx import Document
import openai
from PyQt5.QtCore import QProcess
from concurrent.futures import ThreadPoolExecutor
from fake_useragent import UserAgent
from requests.exceptions import RequestException
from PyQt5.QtWidgets import QApplication, QWidget, QTreeWidget, QTreeWidgetItem, QMessageBox
from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import QVBoxLayout, QHBoxLayout, QGroupBox
from PyQt5.QtWidgets import QToolButton
from PyQt5.QtWidgets import QSpinBox
from openai.error import RateLimitError
import re
import time
import requests
from PyQt5.QtCore import QThread
from PyQt5.QtCore import QRunnable, pyqtSlot, QObject, pyqtSignal
from PyQt5.QtCore import QThread, pyqtSignal, QTimer, QThreadPool
from concurrent.futures import ThreadPoolExecutor
from PyQt5.QtCore import QRunnable, QThreadPool, pyqtSlot
from PyQt5.QtWidgets import QApplication, QMainWindow, QPushButton
import sys
import time
import json
from PyQt5.QtWidgets import QInputDialog

    
class ClickableQTextBrowser(QTextBrowser):
    def __init__(self, parent=None):
        super(ClickableQTextBrowser, self).__init__(parent)
        self.setOpenExternalLinks(False)  # disable the default link handling behavior
        self.anchorClicked.connect(self.on_anchor_clicked)

    def on_anchor_clicked(self, url: QUrl):
        QDesktopServices.openUrl(url)  # open the URL in the default web browser

# Initialize the hyphenator with your desired language
hyphenator = pyphen.Pyphen(lang='en')

# Use the hyphenator to hyphenate a word
hyphenated_word = hyphenator.inserted('hyphenation')

print(hyphenated_word)  # Output: hy-phen-a-tion

class CustomTextEdit(QTextEdit):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.init_scrollbars()

    def init_scrollbars(self):
        style = '''
            QScrollBar:vertical {
                width: 12px;
                background: #f5f5f5;
                margin: 0;
            }
            QScrollBar::handle:vertical {
                background: #d3d3d3;
                min-height: 30px;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0;
                width: 0;
            }
        '''

        self.verticalScrollBar().setStyleSheet(style)
        self.horizontalScrollBar().setStyleSheet(style)
        
class CustomTextEdit(QTextEdit):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.hyphenator = pyphen.Pyphen(lang='en')

    def insertFromMimeData(self, source):
        text = source.text()
        hyphenated_text = self.hyphenate_text(text)
        self.insertPlainText(hyphenated_text)

    def hyphenate_text(self, text):
        words = text.split()
        hyphenated_words = [self.hyphenator.inserted(w, '-') for w in words]
        return ' '.join(hyphenated_words)


class Worker(QRunnable):
    def __init__(self, func, *args, **kwargs):
        super(Worker, self).__init__()
        self.func = func
        self.args = args
        self.kwargs = kwargs
        self.signals = WorkerSignals()

    @pyqtSlot()
    def run(self):
        response = self.func(*self.args, **self.kwargs)
        self.signals.finished.emit(response)


class WorkerSignals(QObject):
    finished = pyqtSignal(str)


class ChatbotAssistantWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.init_ui()
        self.threadpool = QThreadPool()
    def init_ui(self):
        self.layout = QVBoxLayout()
        self.threadpool = QThreadPool()
        self.conversation_box = ClickableQTextBrowser()
        self.conversation_box.setAcceptRichText(True)  # allow the QTextBrowser to display rich text
        # Create a text box to display the conversation
        self.conversation_box = QTextBrowser(self)
        self.conversation_box.setReadOnly(True)
        self.conversation_box.setHtml('Assist: Hello! I am Admin Assist. How can I help you today?')
        self.conversation_box.setVerticalScrollBarPolicy
        (Qt.ScrollBarAlwaysOn)  # Add a scrollbar to the conversation box
        # Enable clickable URL links and open them in the default desktop browser
        self.conversation_box.setOpenExternalLinks(True)
        # Enable clickable email links and open them in the default desktop email client
        self.conversation_box.setOpenLinks(True)
        # enable bullet points in the conversation box
        self.conversation_box.setAcceptRichText(True)
        # Set the font size of the conversation box
        self.conversation_box.setFontPointSize(12)
        # Set the font family of the conversation box
        self.conversation_box.setFontFamily("Arial")
        # Set the font color of the conversation box
        # Set the background color of the conversation box
        self.conversation_box.setStyleSheet("background-color: #f3f3f3;")
        # Set the text alignment of the conversation box 
        self.conversation_box.setAlignment(Qt.AlignLeft)
        
        self.layout.addWidget(self.conversation_box)
        
        self.prompt_label = QLabel("Enter your question:")
        self.layout.addWidget(self.prompt_label)

        self.prompt_input = QPlainTextEdit()
        self.prompt_input.setFixedHeight(50)
        self.layout.addWidget(self.prompt_input)

        self.buttons_layout = QHBoxLayout()

        self.send_button = QPushButton("Send")
        self.send_button.clicked.connect(self.send_message)
        self.buttons_layout.addWidget(self.send_button)

        self.clear_conversation_button = QPushButton("Clear Conversation")
        self.clear_conversation_button.clicked.connect(self.clear_conversation)
        self.buttons_layout.addWidget(self.clear_conversation_button)

        self.clear_prompt_button = QPushButton("Clear Prompt")
        self.clear_prompt_button.clicked.connect(self.clear_prompt)
        self.buttons_layout.addWidget(self.clear_prompt_button)
        
        self.web_search_toggle = QCheckBox("Web Search")
        self.buttons_layout.addWidget(self.web_search_toggle)

        self.layout.addLayout(self.buttons_layout)
        self.setLayout(self.layout)

    def send_message(self):
        user_message = self.prompt_input.toPlainText().strip()
        if user_message:
            self.conversation_box.append(f"<b>You:</b> {user_message}")

            if self.web_search_toggle.isChecked():
                worker = Worker(self.internet_search1, user_message)
            else:
                worker = Worker(openai_chat2, user_message)

            worker.signals.finished.connect(self.update_ui)
            self.threadpool.start(worker)

            self.prompt_input.clear()

        
    def update_ui(self, generated_response):
        self.conversation_box.append(f"<b>Assistant:</b> {generated_response}")
        self.conversation_box.moveCursor(QTextCursor.End)
        self.conversation_box.insertPlainText("\n")  # Adds a newline after the message
        
    def clear_conversation(self):
            self.conversation_box.clear()

    def clear_prompt(self):
        self.prompt_input.clear()
    
    @pyqtSlot()
    def internet_search(self, query):
        sources = [
            {
                "endpoint": "https://en.wikipedia.org/w/api.php",
                "params": {
                    "action": "query",
                    "list": "search",
                    "srsearch": query,
                    "format": "json",
                },
                "response_handler": self._handle_wikipedia_response
            }
        ]
        
        headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3',
            }
        
        items = []
        for source in sources:
            response = requests.get(source["endpoint"], params=source["params"], headers=headers)
            response.raise_for_status()
            items.extend(source["response_handler"](response))

        search_results_text = "\n".join(f"{item['title']}\n{item['url']}" for item in items)
        
        prompt = f"Given the web results, answer the query '{query}' and list citations using sources from the web results. Make HTML links clickable"
        conversation_history.append({"role": "user", "content":f"{prompt}\n\nSearch results:\n{search_results_text}"})
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=conversation_history
            )
            message = response.choices[0].message.content
        except openai.error.APIConnectionError:
            error_message = "Error communicating with server on the internet. Please check your internet connection and try again."
            self.end_task()
            return error_message
        except openai.error.InvalidRequestError:
            error_message = "There was an error processing your request. Please try again later."
            self.end_task()
            return error_message
        except openai.error.RateLimitError as e:
            error_message = f"That model is currently overloaded with other requests. You can retry your request, or contact us through our help center at help.openai.com if the error persists. (Please include the request ID {e.request_id} in your message.)"
            self.end_task()
            return error_message
        return message.strip()
    
    def _get_url_content(self, url, params):
        ua = UserAgent()
        headers = {"User-Agent": ua.random}
        try:
            response = requests.get(url, params=params, headers=headers)
            response.raise_for_status()
            return response, None
        except requests.RequestException as e:
            error_message = f"An error occurred while fetching {url}: {e}"
            return None, error_message


    def _handle_duckduckgo_response(self, response):
        soup = BeautifulSoup(response.text, "html.parser")  # <-- Note the use of `.text` here.
        results = soup.select(".result__url")[:3]
        items = []
        for result in results:
            title = result.text
            url = result.get("href")
            items.append({"title": title, "url": url})  # Append a dictionary instead of a string
        return items


    def _handle_wikipedia_response(self, response):
        results = response.json()["query"]["search"][:1]
        items = []
        for result in results:
            title = result["title"]
            url = f"https://en.wikipedia.org/wiki/{title.replace(' ', '_')}"
            items.append({"title": title, "url": url})
        return items

    @pyqtSlot()
    def internet_search1(self, query):
        sources = [
            {
                "endpoint": "https://en.wikipedia.org/w/api.php",
                "params": {
                    "action": "query",
                    "list": "search",
                    "srsearch": query,
                    "format": "json",
                },
                "response_handler": self._handle_wikipedia_response
            },
            {
                "endpoint": "https://duckduckgo.com/html/",
                "params": {
                    "q": query,
                },
                "response_handler": self._handle_duckduckgo_response
            }
        ]

        items = []
        with ThreadPoolExecutor() as executor:
            for source in sources:
                future_content = executor.submit(self._get_url_content, source["endpoint"], source["params"])
                response, error_message = future_content.result()
                if response is None:
                    print(error_message)
                    self.conversation_box.append(f"<div class='right-bubble'><span style='font-weight:bold;color:red;'>Admin Assist Error: </span>An error occurred while searching the internet could not gather all relevant data</div>")
                    continue
                items.extend(source["response_handler"](response))
                time.sleep(1)  # Wait for 1 second
            
        # Call web_scrape_tool on each URL
        scraped_data_results = []
        for item in items:
            url = item['url']
            scraped_data = self.scrape_data(url)
            scraped_data_results.append(scraped_data)

        search_results_text = "\n".join(f"{item['title']}\n{item['url']}" for item in items)
        
        
        prompt = f"Given the web results, answer the query '{query}' and list citations. Make HTML links clickable"
        conversation_history.append({"role": "user", "content":f"{prompt}\n\nSearch results:\n{search_results_text}"})
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=conversation_history
            )
            message = response.choices[0].message.content
        except openai.error.APIConnectionError:
            error_message = "Error communicating with server on the internet. Please check your internet connection and try again."
            return error_message
        except openai.error.InvalidRequestError:
            error_message = "There was an error processing your request. Please try again later."
            return error_message
        except openai.error.RateLimitError as e:
            error_message = f"That model is currently overloaded with other requests. You can retry your request, or contact us through our help center at help.openai.com if the error persists. (Please include the request ID {e.request_id} in your message.)"
            return error_message
        return message.strip()

    def scrape_data(self, url):
        with ThreadPoolExecutor() as executor:
            try:
                future_content = executor.submit(fetch_url_content, url)
                content = future_content.result()

                if content is None:
                    return "Failed to fetch the URL content."

                future_text = executor.submit(extract_text, content)
                text = future_text.result()

                print("\033[90m\033[3m"+"Scrape completed. Length:" +str(len(text))+".Now extracting relevant info..."+"...\033[0m")
                objective = "Extract relevant information from the plain text"
                task = "Produce data for a search result"  # Replace with your current task
                future_info = executor.submit(extract_relevant_info, objective, text[0:5000], task)
                info = future_info.result()

                future_links = executor.submit(extract_links, content)
                links = future_links.result()

                result = info
                return result
            except Exception as e:
                print(f"An error occurred while scraping data from {url}: {e}")
                self.conversation_box.append(f"<div class='right-bubble'><span style='font-weight:bold;color:red;'>Admin Assist Error: </span>An error occurred while searching the internet and interacting with {url}</div>")

    @pyqtSlot()
    def internet_search2(self, query):
        sources = [
            {
                "endpoint": "https://en.wikipedia.org/w/api.php",
                "params": {
                    "action": "query",
                    "list": "search",
                    "srsearch": query,
                    "format": "json",
                },
                "response_handler": self._handle_wikipedia_response
            },
        ]
        
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/90.0.4430.93 Safari/537.36"
        }
        
        items = []
        for source in sources:
            try:
                response = requests.get(source["endpoint"], params=source["params"], headers=headers)
                response.raise_for_status()
            except RequestException as e:
                print(f"An error occurred while fetching {source['endpoint']}: {e}")
                self.conversation_box.append(f"<div class='right-bubble'><span style='font-weight:bold;color:red;'>Admin Assist Error: </span>An error occurred while searching the internet</div>")
                continue
            items.extend(source["response_handler"](response))
            time.sleep(1)  # Wait for 1 second

        # Call web_scrape_tool on each URL
        scraped_data_results = []
        for item in items:
            url = item['url']
            scraped_data = self.web_scrape_tool(url, self.task, self.objective)
            scraped_data_results.append(scraped_data)

        search_results_text = "\n".join(f"{item['title']}\n{item['url']}" for item in items)
            
        prompt = f"Using the following data extracted from search results {items}, answer the query '{query}' and list citations. Make HTML links clickable"
        conversation_history.append({"role": "user", "content":f"{prompt}\n\ndata:\n{search_results_text}"})
        self.conversation_box.append(f"<div class='right-bubble'><span style='font-weight:bold;color:green;'>Admin Assist: Completing background research </span>Search results:\n{search_results_text}</div>")
        
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=conversation_history
            )
            message = response.choices[0].message.content
        except openai.error.APIConnectionError:
            error_message = "Error communicating with server on the internet. Please check your internet connection and try again."
            self.end_task()
            return error_message
        except openai.error.InvalidRequestError:
            error_message = "There was an error processing your request. Please try again later."
            self.end_task()
            return error_message
        except openai.error.RateLimitError as e:
            error_message = f"That model is currently overloaded with other requests. You can retry your request, or contact us through our help center at help.openai.com if the error persists. (Please include the request ID {e.request_id} in your message.)"
            self.end_task()
            return error_message
        
        return message.strip(), scraped_data_results

    def analyze_url_send_message(self, url):
        content = extract_page_content(url)
        truncated_content = truncate_text(content, 400)  # You can change 400 to the number of tokens you want
        return truncated_content
    
    
    def web_scrape_tool(self, url: str, task:str, objective:str):
        with ThreadPoolExecutor() as executor:
            future_content = executor.submit(fetch_url_content, url)
            content = future_content.result()

            if content is None:
                return "Failed to fetch the URL content."

            future_text = executor.submit(extract_text, content)
            text = future_text.result()

            print("\033[90m\033[3m"+"Scrape completed. Length:" +str(len(text))+".Now extracting relevant info..."+"...\033[0m")
            self.conversation_box.append(f"<div class='right-bubble'><span style='font-weight:bold;color:green;'>Admin Assist: </span>Scrape completed. Length: {str(len(text))}. Now extracting relevant info...</div>")
            
            future_info = executor.submit(extract_relevant_info, objective, text[0:5000], task)
            info = future_info.result()

            future_links = executor.submit(extract_links, content)
            links = future_links.result()

            result = info
            return result
        
class PaginatedTextEdit(QScrollArea):
    def __init__(self, parent=None):
        super().__init__(parent)

        self.text_editor = CustomTextEdit(self)
        self.setWidget(self.text_editor)
        self.setWidgetResizable(True)
        self.text_editor.setWordWrapMode(QTextOption.WordWrap)

        tab_size = 4
        text_option = QTextOption()
        text_option.setTabStopDistance(tab_size * QFontMetricsF(self.text_editor.font()).horizontalAdvance(' '))
        self.text_editor.document().setDefaultTextOption(text_option)

    def paintEvent(self, event):
        painter = QPainter(self.viewport())
        pixmap = QPixmap(self.viewport().size())
        pixmap.fill(Qt.white)

        temp_painter = QPainter(pixmap)
        temp_painter.setRenderHint(QPainter.Antialiasing)

        printer = QPrinter(QPrinter.ScreenResolution)
        printer.setPageSize(QPrinter.A4)
        printer.setPageMargins(10, 10, 10, 10, QPrinter.Millimeter)

        self.text_editor.document().drawContents(temp_painter, printer.pageRect())

        temp_painter.end()
        painter.drawPixmap(0, 0, pixmap)

class WordEditor(QMainWindow):
    error_occurred = pyqtSignal(str)
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Word Editor with Assistant")
        self.setGeometry(100, 100, 800, 600)
        self.current_file = None
        self.load_openai_key()

        self.text_editor = CustomTextEdit(self)
        self.text_editor.textChanged.connect(self.update_title)
        self.setCentralWidget(self.text_editor)
        self.text_editor.setWordWrapMode(QTextOption.WordWrap)

        self.chatbot_assistant = ChatbotAssistantWidget()
        self.dock_widget = QDockWidget("Assistant", self)
        self.dock_widget.setWidget(self.chatbot_assistant)
        self.addDockWidget(Qt.RightDockWidgetArea, self.dock_widget)
        self.dock_widget.hide()
        tab_size = 4
        text_option = QTextOption()
        text_option.setTabStopDistance(tab_size * QFontMetricsF(self.text_editor.font()).horizontalAdvance(' '))
        self.text_editor.document().setDefaultTextOption(text_option)

        # Create the toolbar
        self.init_toolbar()

        # Create the status bar
        self.init_statusbar()

        self.create_menu_bar()

        # Set a modern and appealing stylesheet
        self.set_app_stylesheet()

    def init_toolbar(self):
        toolbar = self.addToolBar("Toolbar")
        toolbar.setMovable(False)

        # Create new document action
        new_action = QAction(QIcon.fromTheme("document-new"), "New", self)
        new_action.setShortcut("Ctrl+N")
        new_action.setToolTip("Create a new document")
        new_action.triggered.connect(self.new_file)
        toolbar.addAction(new_action)

        # Create open document action
        open_action = QAction(QIcon.fromTheme("document-open"), "Open", self)
        open_action.setShortcut("Ctrl+O")
        open_action.setToolTip("Open an existing document")
        open_action.triggered.connect(self.open_file)
        toolbar.addAction(open_action)

        # Create save document action
        save_action = QAction(QIcon.fromTheme("document-save"), "Save", self)
        save_action.setShortcut("Ctrl+S")
        save_action.setToolTip("Save the current document")
        save_action.triggered.connect(self.save_document)
        toolbar.addAction(save_action)
        

        # Create save document as action
        save_as_action = QAction(QIcon.fromTheme("document-save-as"), "Save As", self)
        save_as_action.setShortcut("Ctrl+Shift+S")
        save_as_action.setToolTip("Save the current document with a new name")
        save_as_action.triggered.connect(self.save_document_as)
        toolbar.addAction(save_as_action)

        toolbar.addSeparator()

        # Create undo action
        undo_action = QAction(QIcon.fromTheme("edit-undo"), "Undo", self)
        undo_action.setShortcut("Ctrl+Z")
        undo_action.setToolTip("Undo the last action")
        undo_action.triggered.connect(self.text_editor.undo)
        toolbar.addAction(undo_action)

        # Create redo action
        redo_action = QAction(QIcon.fromTheme("edit-redo"), "Redo", self)
        redo_action.setShortcut("Ctrl+Shift+Z")
        redo_action.setToolTip("Redo the last undone action")
        redo_action.triggered.connect(self.text_editor.redo)
        toolbar.addAction(redo_action)

        toolbar.addSeparator()

        # Create font selection action
        font_action = QAction(QIcon.fromTheme("preferences-desktop-font"), "Font", self)
        font_action.setShortcut("Ctrl+T")
        font_action.setToolTip("Select the font for the text")
        font_action.triggered.connect(self.select_font)
        toolbar.addAction(font_action)
        
        # Create font size selection action
        font_size_action = QAction(QIcon.fromTheme("format-font-size"), "Font Size", self)
        font_size_action.setToolTip("Select the font size for the text")
        font_size_action.triggered.connect(self.select_font_size)
        toolbar.addAction(font_size_action)

        # Create insert image action
        insert_image_action = QAction(QIcon.fromTheme("insert-image"), "Insert Image", self)
        insert_image_action.setShortcut("Ctrl+I")
        insert_image_action.setToolTip("Insert an image from your local drive")
        insert_image_action.triggered.connect(self.insert_image)
        toolbar.addAction(insert_image_action)
        
        toolbar.addSeparator()

        # Create paragraph alignment actions
        align_left_action = QAction(QIcon.fromTheme("format-justify-left"), "Align Left", self)
        align_left_action.setToolTip("Align text to the left")
        align_left_action.triggered.connect(lambda: self.set_paragraph_alignment(Qt.AlignLeft))
        toolbar.addAction(align_left_action)

        align_center_action = QAction(QIcon.fromTheme("format-justify-center"), "Align Center", self)
        align_center_action.setToolTip("Center text")
        align_center_action.triggered.connect(lambda: self.set_paragraph_alignment(Qt.AlignCenter))
        toolbar.addAction(align_center_action)

        align_right_action = QAction(QIcon.fromTheme("format-justify-right"), "Align Right", self)
        align_right_action.setToolTip("Align text to the right")
        align_right_action.triggered.connect(lambda: self.set_paragraph_alignment(Qt.AlignRight))
        toolbar.addAction(align_right_action)
        
        # Create bold text action
        bold_action = QAction(QIcon.fromTheme("format-text-bold"), "Bold", self)
        bold_action.setShortcut("Ctrl+B")
        bold_action.setCheckable(True)
        bold_action.setToolTip("Toggle bold text")
        bold_action.triggered.connect(self.toggle_bold_text)
        toolbar.addAction(bold_action)

        # Create italics text action
        italic_action = QAction(QIcon.fromTheme("format-text-italic"), "Italic", self)
        italic_action.setShortcut("Ctrl+I")
        italic_action.setCheckable(True)
        italic_action.setToolTip("Toggle italic text")
        italic_action.triggered.connect(self.toggle_italic_text)
        toolbar.addAction(italic_action)
        
        # Create underline text action
        underline_action = QAction(QIcon.fromTheme("format-text-underline"), "Underline", self)
        underline_action.setShortcut("Ctrl+U")
        underline_action.setCheckable(True)
        underline_action.setToolTip("Toggle underline text")
        underline_action.triggered.connect(self.toggle_underline_text)
        toolbar.addAction(underline_action)

        # Create strike through text action
        strike_action = QAction(QIcon.fromTheme("format-text-strikethrough"), "Strike Through", self)
        strike_action.setCheckable(True)
        strike_action.setToolTip("Toggle strike through text")
        strike_action.triggered.connect(self.toggle_strike_text)
        toolbar.addAction(strike_action)

        toolbar.addSeparator()

        # Create increase text size action
        increase_size_action = QAction(QIcon.fromTheme("format-font-size-more"), "Increase Text Size", self)
        increase_size_action.setShortcut("Ctrl++")
        increase_size_action.setToolTip("Increase text size")
        increase_size_action.triggered.connect(self.increase_text_size)
        toolbar.addAction(increase_size_action)

        # Create decrease text size action
        decrease_size_action = QAction(QIcon.fromTheme("format-font-size-less"), "Decrease Text Size", self)
        decrease_size_action.setShortcut("Ctrl+-")
        decrease_size_action.setToolTip("Decrease text size")
        decrease_size_action.triggered.connect(self.decrease_text_size)
        toolbar.addAction(decrease_size_action)

        toolbar.addSeparator()

        # Create generate image action
        generate_image_action = QAction(QIcon.fromTheme("image-x-generic"), "Generate Image", self)
        generate_image_action.setShortcut("Ctrl+G")
        generate_image_action.setToolTip("Generate an image from a description")
        generate_image_action.triggered.connect(self.generate_image_dialog)
        toolbar.addAction(generate_image_action)

        toolbar.addSeparator()

        # Create spell check action
        spell_check_action = QAction(QIcon.fromTheme("tools-check-spelling"), "Spell Check", self)
        spell_check_action.setShortcut("F7")
        spell_check_action.setToolTip("Check the spelling of the text")
        spell_check_action.triggered.connect(self.spell_check)
        toolbar.addAction(spell_check_action)
        
        toolbar.addSeparator()
        
        # Create zoom in action
        zoom_in_action = QAction(QIcon.fromTheme("zoom-in"), "Zoom In", self)
        zoom_in_action.setShortcut("Ctrl++")
        zoom_in_action.setToolTip("Zoom in")
        zoom_in_action.triggered.connect(self.zoom_in)
        toolbar.addAction(zoom_in_action)

        # Create zoom out action
        zoom_out_action = QAction(QIcon.fromTheme("zoom-out"), "Zoom Out", self)
        zoom_out_action.setShortcut("Ctrl+-")
        zoom_out_action.setToolTip("Zoom out")
        zoom_out_action.triggered.connect(self.zoom_out)
        toolbar.addAction(zoom_out_action)
    
    def zoom_in(self):
        self.text_editor.zoomIn()

    def zoom_out(self):
        self.text_editor.zoomOut()


    def generate_image_dialog(self):
        description, ok = QInputDialog.getText(self, "Image Description", "Enter a description of the image:")
        if ok and description:
            pixmap = self.generate_image(description)
            
            # Save the QPixmap to a temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            pixmap.save(temp_file.name)
            temp_file.close()

            # Insert the image into the text editor
            self.text_editor.insert_image(temp_file.name)
            
            # Display the image in the ImageWindow
            image_window = ImageWindow(pixmap, self)
            image_window.exec_()
    
    def insert_image(self):
        options = QFileDialog.Options()
        options |= QFileDialog.ReadOnly
        file_name, _ = QFileDialog.getOpenFileName(self, "Insert Image", "",
                                                "Images (*.png *.xpm *.jpg *.gif *.svg);;All Files (*)", options=options)
        if file_name:
            cursor = self.text_editor.textCursor()
            cursor.insertHtml(f'<img src="{file_name}" />')



    def generate_image(self, prompt, variation=False):
        if not variation:
            response = openai.Image.create(
                prompt=prompt,
                n=1,
                size="1024x1024"
            )
        else:
            response = openai.Image.create_variation(
                image=open("last_generated_image.png", "rb"),
                n=1,
                size="1024x1024"
            )

        image_url = response['data'][0]['url']
        img_data = requests.get(image_url).content

        with open("last_generated_image.png", 'wb') as f:
            f.write(img_data)

        qimg = QImage()
        qimg.loadFromData(BytesIO(img_data).read())
        pixmap = QPixmap.fromImage(qimg)

        return pixmap
    
    def update_title(self):
        self.setWindowTitle(f"{self.document_title} - PyText")
    
    def document_title(self):
        self.update_title()
    
    def toggle_bold_text(self, checked):
        fmt = QTextCharFormat()
        fmt.setFontWeight(QFont.Bold if checked else QFont.Normal)
        self.text_editor.mergeCurrentCharFormat(fmt)

    def toggle_italic_text(self, checked):
        fmt = QTextCharFormat()
        fmt.setFontItalic(checked)
        self.text_editor.mergeCurrentCharFormat(fmt)
    
    def init_statusbar(self):
        self.statusbar = QStatusBar()
        self.setStatusBar(self.statusbar)
    
    def toggle_underline_text(self, checked):
        fmt = QTextCharFormat()
        fmt.setFontUnderline(checked)
        self.text_editor.mergeCurrentCharFormat(fmt)

    def toggle_strike_text(self, checked):
        fmt = QTextCharFormat()
        fmt.setFontStrikeOut(checked)
        self.text_editor.mergeCurrentCharFormat(fmt)

    def increase_text_size(self):
        fmt = QTextCharFormat()
        font = self.text_editor.currentCharFormat().font()
        font.setPointSize(font.pointSize() + 1)
        fmt.setFont(font)
        self.text_editor.mergeCurrentCharFormat(fmt)

    def decrease_text_size(self):
        fmt = QTextCharFormat()
        font = self.text_editor.currentCharFormat().font()
        font.setPointSize(max(font.pointSize() - 1, 1))
        fmt.setFont(font)
        self.text_editor.mergeCurrentCharFormat(fmt)

    
    def update_word_count(self, word_count):
        self.word_count_label.setText(f"Word Count: {word_count}")
            
    def create_menu_bar(self):
        menu_bar = self.menuBar()

        file_menu = menu_bar.addMenu("File")

        new_file_action = QAction("New", self)
        new_file_action.triggered.connect(self.new_file)
        file_menu.addAction(new_file_action)

        open_file_action = QAction("Open", self)
        open_file_action.triggered.connect(self.open_file)
        file_menu.addAction(open_file_action)

        save_file_action = QAction("Save", self)
        save_file_action.triggered.connect(self.save_file)
        file_menu.addAction(save_file_action)

        save_file_as_action = QAction("Save As", self)
        save_file_as_action.triggered.connect(self.save_document_as)
        file_menu.addAction(save_file_as_action)

        exit_action = QAction("Exit", self)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        font_action = QAction("Font", self)
        font_action.triggered.connect(self.select_font)
        file_menu.addAction(font_action)
        
        settings_action = QAction("Settings", self)
        settings_action.triggered.connect(self.open_settings_dialog)
        file_menu.addAction(settings_action)

        # Add the Find and Replace action
        find_replace_action = QAction("Find and Replace", self)
        find_replace_action.triggered.connect(self.open_find_replace_dialog)
        file_menu.addAction(find_replace_action)
        
        api_action = QAction("API", self)
        api_action.triggered.connect(self.open_api_dialog)
        file_menu.addAction(api_action)

    def load_openai_key(self):
        if os.path.exists("api_key.json"):
            with open("api_key.json", "r") as file:
                data = json.load(file)
                if "api_key" in data:
                    self.setup_openai(data["api_key"])
                else:
                    self.open_api_dialog()
        else:
            self.open_api_dialog()
    
    def open_api_dialog(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("OpenAI API Key")
        layout = QVBoxLayout(dialog)

        label = QLabel("Enter your OpenAI API Key:", dialog)
        layout.addWidget(label)

        line_edit = QLineEdit(dialog)
        if os.path.exists("api_key.json"):
            with open("api_key.json", "r") as file:
                data = json.load(file)
                if "api_key" in data:
                    line_edit.setText(data["api_key"])
        layout.addWidget(line_edit)

        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, dialog)
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)
        layout.addWidget(button_box)

        delete_button = QPushButton("Delete API Key", dialog)
        delete_button.clicked.connect(self.delete_api_key)
        layout.addWidget(delete_button)

        dialog.setLayout(layout)

        if dialog.exec() == QDialog.Accepted:
            api_key = line_edit.text()
            if api_key:
                with open("api_key.json", "w") as file:
                    json.dump({"api_key": api_key}, file)
                self.setup_openai(api_key)

    def delete_api_key(self):
        if os.path.exists("api_key.json"):
            msg_box = QMessageBox(self)
            msg_box.setIcon(QMessageBox.Warning)
            msg_box.setWindowTitle("Delete API Key")
            msg_box.setText("Are you sure you want to delete the API key?")
            msg_box.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
            msg_box.setDefaultButton(QMessageBox.No)

            response = msg_box.exec()

            if response == QMessageBox.Yes:
                os.remove("api_key.json")
                QMessageBox.information(self, "Delete API Key", "API key deleted successfully.")
        else:
            QMessageBox.warning(self, "Delete API Key", "No API key exists.")

    def setup_openai(self, api_key):
        openai.api_key = api_key

    def set_app_stylesheet(self):
        style_sheet = """
        QMainWindow {
            background-color: #f0f0f0;
        }
        
        QDockWidget {
            font-size: 12px;
        }

        QTextEdit {
            background-color: #ffffff;
            color: #333333;
            font-size: 14px;
            font-family: 'Calibri';
            selection-background-color: #3399ff;
            selection-color: #ffffff;
            line-height: 1.5;
            padding: 10px;
            border: 1px solid #d3d3d3;
            border-radius: 3px;
        }
        
        QMenuBar {
            background-color: #ffffff;
        }

        QMenuBar::item {
            background-color: #ffffff;
        }

        QMenu {
            background-color: #ffffff;
        }

        QMenu::item {
            background-color: #ffffff;
        }

        QMenu::item:selected {
            background-color: #f0f0f0;
        }

        QToolBar {
            background-color: #ffffff;
            border: none;
        }

        QStatusBar {
            background-color: #ffffff;
        }
        """
        self.setStyleSheet(style_sheet)

    def init_statusbar(self):
        self.statusbar = QStatusBar()
        self.setStatusBar(self.statusbar)
        self.word_count_label = QLabel("Word Count: 0")
        self.statusbar.addPermanentWidget(self.word_count_label)
        
    def open_settings_dialog(self):
        settings_dialog = SettingsDialog(self)
        result = settings_dialog.exec_()

        if result == QDialog.Accepted:
            settings_dialog.font_size()

            # Reload your application's settings here
            
    def select_font_size(self):
        font, ok = QFontDialog.getFont(self.text_editor.font(), self, "Select Font")
        if ok:
            self.text_editor.setFont(font)

    def set_paragraph_alignment(self, alignment):
        cursor = self.text_editor.textCursor()
        block_format = QTextBlockFormat()
        block_format.setAlignment(alignment)
        cursor.mergeBlockFormat(block_format)
        self.text_editor.setTextCursor(cursor)

    def spell_check(self):
        spell = SpellChecker()
        text = self.text_editor.toPlainText()
        words = text.split()
        misspelled = spell.unknown(words)

        if not misspelled:
            QMessageBox.information(self, "Spell Check", "No spelling errors found.")
            return

        for word in misspelled:
            suggestions = spell.candidates(word)
            if suggestions:
                suggested_word = list(suggestions)[0]
                text = text.replace(word, suggested_word)

        self.text_editor.setPlainText(text)
        QMessageBox.information(self, "Spell Check", "Spelling errors have been corrected.")
        
    def open_find_replace_dialog(self):
        find_replace_dialog = FindReplaceDialog(self)
        find_replace_dialog.exec_()

    def new_file(self):
        # Check if the text editor is modified and not empty
        if self.text_editor.document().isModified() and not self.text_editor.toPlainText().strip() == "":
            msg_box = QMessageBox(self)
            msg_box.setIcon(QMessageBox.Warning)
            msg_box.setWindowTitle("New Document")
            msg_box.setText("You have unsaved changes. Do you want to save your current document?")
            msg_box.setStandardButtons(QMessageBox.Save | QMessageBox.Discard | QMessageBox.Cancel)
            msg_box.setDefaultButton(QMessageBox.Save)

            response = msg_box.exec()

            # Save the current document
            if response == QMessageBox.Save:
                self.save_document()

            # Discard changes and create a new document
            elif response == QMessageBox.Discard:
                self.text_editor.clear()
        else:
            # Create a new document if the current document is empty or unchanged
            self.text_editor.clear()

    def open_file(self):
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getOpenFileName(self, "Open Document", "", "Word Files (*.docx);;Text Files (*.txt);;All Files (*)", options=options)
        if file_name:
            if file_name.endswith(".docx"):
                doc = docx.Document(file_name)
                full_text = []
                for paragraph in doc.paragraphs:
                    full_text.append(paragraph.text)
                self.text_editor.setPlainText("\n".join(full_text))
            else:
                file = QFile(file_name)
                if file.open(QIODevice.ReadOnly | QIODevice.Text):
                    stream = QTextStream(file)
                    self.text_editor.setPlainText(stream.readAll())
                    file.close()
        if file_name:
            self.current_file = file_name

    def save_file(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        file_name, _ = QFileDialog.getSaveFileName(self,"Save File", "", "Text Files (.txt);;All Files ()", options=options)
        if file_name:
            with open(file_name, 'w') as file:
                file.write(self.text_editor.toPlainText())

    def save_document(self):
        if self.current_file:
            self.save_to_file(self.current_file)
        else:
            self.save_document_as()

    def save_document_as(self):
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getSaveFileName(self, "Save Document As", "", "Word Files (*.docx);;Text Files (*.txt);;All Files (*)", options=options)
        if file_name:
            self.save_to_file(file_name)
            self.current_file = file_name

    def save_to_file(self, file_name):
        if file_name.endswith(".docx"):
            doc = docx.Document()
            paragraphs = self.text_editor.toPlainText().split("\n")
            for paragraph in paragraphs:
                doc.add_paragraph(paragraph)
            doc.save(file_name)
        else:
            file = QFile(file_name)
            if file.open(QIODevice.WriteOnly | QIODevice.Text):
                stream = QTextStream(file)
                stream << self.text_editor.toPlainText()
                file.close()

    def select_font(self):
        font, ok = QFontDialog.getFont(self.text_editor.font(), self)
        if ok:
            self.text_editor.setFont(font)

class CustomTextEdit(QPlainTextEdit):
    textChangedSignal = pyqtSignal(int)
    
    class Worker(QRunnable):
        '''
        Worker thread
        '''
        def __init__(self, fn, *args, **kwargs):
            super(CustomTextEdit.Worker, self).__init__()
            # Store constructor arguments (re-used for processing)
            self.fn = fn
            self.args = args
            self.kwargs = kwargs
            self.signals = CustomTextEdit.Worker.Signals()
            
        @pyqtSlot()
        def run(self):
            '''
            Initialise the runner function with passed args, kwargs.
            '''
            
            result = self.fn(*self.args, **self.kwargs)
            self.signals.result.emit(result)
            
        class Signals(QObject):
            '''
            Defines the signals available from a running worker thread.
            Supported signals are:
            finished
                No data
            error
                `tuple` (exctype, value, traceback.format_exc() )
            result
                `object` data returned from processing, anything
            '''
            result = pyqtSignal(object)
            finished = pyqtSignal()
            error = pyqtSignal(tuple)
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.textChanged.connect(self.update_word_count)
        self.threadpool = QThreadPool()
  

    def contextMenuEvent(self, event):
        menu = QMenu(self)
        self.threadpool = QThreadPool()
        # Add menu options
        analyze_action = menu.addAction("Analyze")
        improve_action = menu.addAction("Improve")
        assess_grammar_action = menu.addAction("Assess Grammar")
        give_ideas_action = menu.addAction("Give Ideas")
        extend_action = menu.addAction("Extend")
        shorten_action = menu.addAction("Shorten")
        spell_check_action = menu.addAction("Check Spelling")
        autocomplete_action = menu.addAction("Autocomplete")
        copy_action = menu.addAction("Copy")
        paste_action = menu.addAction("Paste")
        reference_action = menu.addAction("Reference")
        cut_action = menu.addAction("Cut")

        
        analyze_action.triggered.connect(self.analyze_text)
        improve_action.triggered.connect(self.improve_text)
        assess_grammar_action.triggered.connect(self.assess_grammar)
        give_ideas_action.triggered.connect(self.give_ideas)
        extend_action.triggered.connect(self.extend_text)
        shorten_action.triggered.connect(self.shorten_text)
        spell_check_action.triggered.connect(self.check_spelling)
        autocomplete_action.triggered.connect(self.autocomplete)
        copy_action.triggered.connect(self.copy_text)
        paste_action.triggered.connect(self.paste_text)
        cut_action.triggered.connect(self.cut_text)
        reference_action.triggered.connect(self.generate_reference)

                                             
        menu.exec_(event.globalPos())
        
    def wheelEvent(self, event):
        if event.modifiers() == Qt.ControlModifier:
            if event.angleDelta().y() > 0:
                self.zoomIn()
            else:
                self.zoomOut()
        else:
            super().wheelEvent(event)


    def cut_text(self):
        self.cut()
    
    def copy_text(self):
        self.copy()

    def paste_text(self):
        self.paste()

    def insert_image(self, file_path):
        cursor = self.textCursor()
        cursor.insertHtml(f'<img src="{file_path}" />')

    def autocomplete(self):
        selected_text = self.textCursor().selectedText()
        if selected_text:
            prompt = f"Autocomplete the following text: {selected_text}"
            worker = self.Worker(self.send_prompt, prompt)
            worker.signals.result.connect(self.insert_text)
            self.threadpool.start(worker)

    def send_prompt(self, prompt):
        self.send_prompt_to_chatbot(prompt)
    
    def insert_text(self, text):
        #activate replace_selected_text
        self.replace_selected_text(text)
        
    def check_spelling(self):
        selected_text = self.textCursor().selectedText()
        if selected_text:
            prompt = f"Check the spelling of the following text: {selected_text}"
            worker = self.Worker(self.send_prompt, prompt)
            worker.signals.result.connect(self.insert_text)
            self.threadpool.start(worker)

    def update_word_count(self):
        word_count = len(self.toPlainText().split())
        self.textChangedSignal.emit(word_count)
        
    def analyze_text(self):
        selected_text = self.textCursor().selectedText()
        if selected_text:
            prompt = f"Analyze the following text: {selected_text}"
            self.send_prompt_to_chatbot(prompt)

    def improve_text(self):
        selected_text = self.textCursor().selectedText()
        if selected_text:
            prompt = f"Improve the following text: {selected_text}"
            self.send_prompt_to_chatbot(prompt)

    def assess_grammar(self):
        selected_text = self.textCursor().selectedText()
        if selected_text:
            prompt = f"Assess the grammar of the following text: {selected_text}"
            self.send_prompt_to_chatbot(prompt)

    def give_ideas(self):
        selected_text = self.textCursor().selectedText()
        if selected_text:
            prompt = f"Generate ideas based on the following text: {selected_text}"
            self.send_prompt_to_chatbot(prompt)

    def extend_text(self):
        selected_text = self.textCursor().selectedText()
        if selected_text:
            prompt = f"Extend the following text: {selected_text}"
            worker = self.Worker(self.send_prompt, prompt)
            worker.signals.result.connect(self.insert_text)
            self.threadpool.start(worker)

    def shorten_text(self):
        selected_text = self.textCursor().selectedText()
        if selected_text:
            prompt = f"Shorten the following text: {selected_text}"
            worker = self.Worker(self.send_prompt, prompt)
            worker.signals.result.connect(self.insert_text)
            self.threadpool.start(worker)
    
    def replace_selected_text(self, new_text):
        cursor = self.textCursor()
        original_text = cursor.selectedText()
        cursor.insertText(original_text + new_text)
    
    def send_prompt_to_chatbot(self, prompt):
        worker = self.Worker(self.send_prompt, prompt)
        worker.signals.result.connect(self.sendchat)
        self.threadpool.start(worker)
    
    def sendchat(self, generated_response):
        self.parent().chatbot_assistant.conversation_box.append(
            f"<b>Assistant:</b> {generated_response}"
        )
    
    def send_prompt(self, prompt):
        response = openai_chat(prompt)  # Replace with openai_chat2 if using the second function
        return response

    @pyqtSlot()
    def generate_reference(self):
        selected_text = self.textCursor().selectedText()
        if selected_text:
            query = f"Find reference for '{selected_text}'"
            worker = self.Worker(self.internet_search, query)
            worker.signals.finished.connect(self.insert_text)
            self.threadpool.start(worker)


    
    def insert_text(self, text):
        cursor = self.textCursor()
        
        # Move the cursor to the end of the selected text
        cursor.setPosition(cursor.selectionEnd())

        # Insert a line break and the reference text
        cursor.insertText(f"\n{text}")

        # Move the cursor to the end of the inserted text
        cursor.setPosition(cursor.position())

        self.setTextCursor(cursor)

        

    @pyqtSlot()
    def internet_search(self, query):
        sources = [
            {
                "endpoint": "https://en.wikipedia.org/w/api.php",
                "params": {
                    "action": "query",
                    "list": "search",
                    "srsearch": query,
                    "format": "json",
                },
                "response_handler": self._handle_wikipedia_response
            },
            {
                "endpoint": "https://duckduckgo.com/html/",
                "params": {
                    "q": query,
                },
                "response_handler": self._handle_duckduckgo_response
            }
        ]
        
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/90.0.4430.93 Safari/537.36"
        }
        
        items = []
        for source in sources:
            response = requests.get(source["endpoint"], params=source["params"], headers=headers)
            response.raise_for_status()
            items.extend(source["response_handler"](response))

        search_results_text = "\n".join(f"{item['title']}\n{item['url']}" for item in items)
        truncated_search_results_text = truncate_text(search_results_text, 400)

        prompt = f"Given the web results, answer the query '{query}' and list citations using sources from the web results. Make HTML links clickable"
        conversation_history.append({"role": "user", "content":f"{prompt}\n\nSearch results:\n{truncated_search_results_text}"})
        # The rest of the internet_search function remains unchanged
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=conversation_history
            )
            message = response.choices[0].message.content
        except openai.error.APIConnectionError:
            error_message = "Error communicating with server on the internet. Please check your internet connection and try again."
            # pop up of error message
            error = QMessageBox()
            error.setWindowTitle("Error")
            error.setText(error_message)
            error.setIcon(QMessageBox.Critical)
            error.exec_()
            return error_message
        except openai.error.InvalidRequestError:
            error_message = "There was an error processing your request. Please try again later."
            return error_message
        except openai.error.RateLimitError as e:
            error_message = f"That model is currently overloaded with other requests. You can retry your request, or contact us through our help center at help.openai.com if the error persists. (Please include the request ID {e.request_id} in your message.)"
            return error_message
        return message.strip()


    
    def _handle_duckduckgo_response(self, response):
        soup = BeautifulSoup(response.text, "html.parser")
        results = soup.select(".result__url")[:3]
        items = []
        for result in results:
            title = result.text
            url = result.get("href")
            items.append({"title": title, "url": url})
        return items

    def _handle_wikipedia_response(self, response):
        json_response = response.json()
        if "query" not in json_response:
            print(f"Error: Unexpected response from Wikipedia API: {json_response}")
            return []
        results = json_response["query"]["search"][:1]
        items = []
        for result in results:
            title = result["title"]
            url = f"https://en.wikipedia.org/wiki/{title.replace(' ', '_')}"
            items.append({"title": title, "url": url})
        return items

        
def openai_chat(message):
        global conversation_history
        conversation_history.append({"role": "user", "content": message})
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=conversation_history
            )
            reply = response.choices[0].message.content
            conversation_history.append({"role": "assistant", "content": reply})
            return reply
        except openai.error.APIConnectionError:
            error_message = "Error communicating with server on the internet. Please check your internet connection and try again."
            return error_message
        except openai.error.InvalidRequestError:
            error_message = "There was an error processing your request. Please try again later."
            return error_message
        except openai.error.RateLimitError as e:
            error_message = "That model is currently overloaded with other requests. You can retry your request, or contact us through our help center at help.openai.com if the error persists. (Please include the request ID {e.request_id} in your message.)"
            return error_message

conversation_history = [
    {"role": "system", "content": 'You are an AI writing assistant'},
]

while len(conversation_history) > 4:  # You can adjust this value based on your requirements
    conversation_history.pop(0)  # Remove the oldest message
prev_prompt = ""

messages = []


def openai_chat2(message):
    global conversation_history
    conversation_history.append({"role": "user", "content": message})
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo-16k",  # Change this line to the new model name
            messages=conversation_history
        )
        reply = response.choices[0].message.content
        conversation_history.append({"role": "assistant", "content": reply})
        return reply
    except openai.error.APIConnectionError:
        error_message = "Error communicatcing with server on the internet. Please check your internet connection and try again."
        return error_message
    except openai.error.InvalidRequestError:
        error_message = "There was an error processing your request. Please try again later."
        return error_message
    except openai.error.RateLimitError as e:
        error_message = "That model is currently overloaded with other requests. You can retry your request, or contact us through our help center at help.openai.com if the error persists. (Please include the request ID {e.request_id} in your message.)"
        return error_message

conversation_history = [
    {"role": "system", "content": 'You are an AI writing assistant you can provide personal opinions and support the user '},]

while len(conversation_history) > 10: # You can adjust this value based on your requirements
        conversation_history.pop(0) # Remove the oldest message


# Instantiate GPT2 tokenizer
tokenizer = GPT2Tokenizer.from_pretrained("gpt2")

# Define headers
headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/90.0.4430.93 Safari/537.36"
}

# Define the truncate_text function
def truncate_text(text: str, max_tokens: int) -> str:
    tokenizer = GPT2Tokenizer.from_pretrained("gpt2")
    tokens = tokenizer.encode(text, return_tensors="pt")

    if tokens.shape[1] <= max_tokens:
        return text

    truncated_tokens = tokens[0, :max_tokens].tolist()
    truncated_text = tokenizer.decode(truncated_tokens, skip_special_tokens=True)
    return truncated_text

def extract_page_content(url: str) -> str:
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/90.0.4430.93 Safari/537.36"
    }

    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
    except requests.exceptions.HTTPError as e:
        print(f"Error fetching content for {url}: {e}")
        return "Error fetching content"
        
    soup = BeautifulSoup(response.text, "html.parser")

    # Remove unnecessary elements, such as tables, citations, and external links
    for unwanted_element in soup(["table", "sup", "span"]):
        unwanted_element.decompose()

    # Extract the main content text from the 'mw-content-text' div element
    main_content = soup.find("div", {"id": "mw-content-text"})
    if main_content:
        text = ' '.join(paragraph.text for paragraph in main_content.find_all("p"))
    else:
        text = ""

    return text.strip()
 
headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/94.0.4606.81 Safari/537.36"
}
       
def fetch_url_content(url: str):
    try:
        # Prepend the scheme (https://) to the URL, only if it's not already there
        full_url = url if url.startswith("https://") else "https:" + url
        response = requests.get(full_url, headers=headers, timeout=10)
        response.raise_for_status()
        return response.content
    except requests.exceptions.RequestException as e:
        print(f"Error while fetching the URL {url}: {e}")
        return None

def extract_links(content: str):
    soup = BeautifulSoup(content, "html.parser")
    links = [link.get('href') for link in soup.findAll('a', attrs={'href': re.compile("^https?://")})]
    return links

def extract_text(content: str):
    soup = BeautifulSoup(content, "html.parser")
    text = soup.get_text(strip=True)
    return text if text else "No text found in content."

def extract_relevant_info(objective, large_string, task):
    chunk_size = 3000
    overlap = 500
    notes = ""
    
    for i in range(0, len(large_string), chunk_size - overlap):
        chunk = large_string[i:i + chunk_size]
        
        messages = [
            {"role": "system", "content": f"Objective: {objective}\nCurrent Task:{task}"},
            {"role": "user", "content": f"Analyze the following text and extract information relevant to our objective and current task, and only information relevant to our objective and current task. If there is no relevant information do not say that there is no relevant information related to our objective. ### Then, update or start our notes provided here (keep blank if currently blank): {notes}.### Text to analyze: {chunk}.### Updated Notes:"}
        ]

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=messages,
            max_tokens=800,
            n=1,
            stop="###",
            temperature=0.7,
        )

        notes += response.choices[0].message['content'].strip()+". "
    
    return notes if notes.strip() else "No relevant information found."

settings = QSettings("YourOrganization", "YourApplication")
# Example: Load font size
font_size = settings.value("fontSize", defaultValue=12, type=int)

# Settings Dialog
class SettingsDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()

        # Theme setting
        theme_layout = QHBoxLayout()
        theme_label = QLabel("Theme:")
        self.theme_combo = QComboBox()
        self.theme_combo.addItem("Light")
        self.theme_combo.addItem("Dark")
        theme_layout.addWidget(theme_label)
        theme_layout.addWidget(self.theme_combo)
        layout.addLayout(theme_layout)

        # Default font setting
        font_layout = QHBoxLayout()
        font_label = QLabel("Default Font:")
        self.font_combo = QComboBox()
        self.font_combo.addItem("Arial")
        self.font_combo.addItem("Times New Roman")
        self.font_combo.addItem("Courier New")
        font_layout.addWidget(font_label)
        font_layout.addWidget(self.font_combo)
        layout.addLayout(font_layout)
        
        #font size setting
        font_size_layout = QHBoxLayout()
        font_size_label = QLabel("Font Size:")
        self.font_size_spinbox = QSpinBox()
        self.font_size_spinbox.setRange(8, 24)
        font_size_layout.addWidget(font_size_label)
        font_size_layout.addWidget(self.font_size_spinbox)
        layout.addLayout(font_size_layout)
        

        # Enable/disable spell checking
        self.spell_check_checkbox = QCheckBox("Enable spell checking")
        layout.addWidget(self.spell_check_checkbox)

        # Buttons
        buttons_layout = QHBoxLayout()
        self.save_button = QPushButton("Save")
        self.cancel_button = QPushButton("Cancel")
        buttons_layout.addWidget(self.save_button)
        buttons_layout.addWidget(self.cancel_button)
        layout.addLayout(buttons_layout)

        self.setLayout(layout)

        self.save_button.clicked.connect(self.accept)
        self.cancel_button.clicked.connect(self.reject)
    
    def font_size(self):
        new_fontsize = self.font_size_spinbox.value()
        settings.setValue("fontSize", new_fontsize)

    

# Find and Replace Dialog
class FindReplaceDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("Find and Replace")

        layout = QVBoxLayout()

        # Find input field
        find_layout = QHBoxLayout()
        find_label = QLabel("Find:")
        self.find_input = QLineEdit()
        find_layout.addWidget(find_label)
        find_layout.addWidget(self.find_input)
        layout.addLayout(find_layout)

        # Replace input field
        replace_layout = QHBoxLayout()
        replace_label = QLabel("Replace with:")
        self.replace_input = QLineEdit()
        replace_layout.addWidget(replace_label)
        replace_layout.addWidget(self.replace_input)
        layout.addLayout(replace_layout)

        # Options
        options_layout = QVBoxLayout()
        self.case_sensitive_checkbox = QCheckBox("Case sensitive")
        self.whole_word_checkbox = QCheckBox("Whole words only")
        self.regex_checkbox = QCheckBox("Regular expression")
        options_layout.addWidget(self.case_sensitive_checkbox)
        options_layout.addWidget(self.whole_word_checkbox)
        options_layout.addWidget(self.regex_checkbox)
        layout.addLayout(options_layout)

        # Buttons
        buttons_layout = QHBoxLayout()
        self.find_button = QPushButton("Find")
        self.replace_button = QPushButton("Replace")
        self.replace_all_button = QPushButton("Replace All")
        self.cancel_button = QPushButton("Cancel")
        buttons_layout.addWidget(self.find_button)
        buttons_layout.addWidget(self.replace_button)
        buttons_layout.addWidget(self.replace_all_button)
        buttons_layout.addWidget(self.cancel_button)
        layout.addLayout(buttons_layout)

        self.setLayout(layout)

        self.find_button.clicked.connect(self.find)
        self.replace_button.clicked.connect(self.replace)
        self.replace_all_button.clicked.connect(self.replace_all)
        self.cancel_button.clicked.connect(self.reject)

    def find(self):
        find_text = self.find_input.text()
        if not find_text:
            return

        options = QTextDocument.FindFlags()

        if self.case_sensitive_checkbox.isChecked():
            options |= QTextDocument.FindCaseSensitively

        if self.whole_word_checkbox.isChecked():
            options |= QTextDocument.FindWholeWords

        if self.regex_checkbox.isChecked():
            raise NotImplementedError("Regular expression search is not implemented.")

        found = self.parent.find(find_text, options)

        if not found:
            self.parent.moveCursor(QTextCursor.Start)

    def replace(self):
        cursor = self.parent.textCursor()
        if cursor.hasSelection():
            cursor.insertText(self.replace_input.text())

        self.find()

    def replace_all(self):
        self.parent.moveCursor(QTextCursor.Start)
        find_text = self.find_input.text()
        replace_text = self.replace_input.text()

        options = QTextDocument.FindFlags()

        if self.case_sensitive_checkbox.isChecked():
            options |= QTextDocument.FindCaseSensitively

        if self.whole_word_checkbox.isChecked():
            options |= QTextDocument.FindWholeWords

        if self.regex_checkbox.isChecked():
            raise NotImplementedError("Regular expression search is not implemented.")

        while self.parent.find(find_text, options):
            cursor = self.parent.textCursor()
            if cursor.hasSelection():
                cursor.insertText(replace_text)

class ImageWindow(QDialog):
    def __init__(self, pixmap, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Generated Image")

        self.original_pixmap = pixmap
        self.scaled_pixmap = self.original_pixmap.scaled(1000, 1200, Qt.KeepAspectRatio)

        self.image_label = QLabel(self)
        self.image_label.setPixmap(self.scaled_pixmap)
        self.image_label.setSizePolicy(QSizePolicy.Ignored, QSizePolicy.Ignored)
        self.image_label.setScaledContents(True)

        # Make the image clickable
        self.image_label.mousePressEvent = self.open_save_image_dialog

        # Create buttons
        self.save_button = QPushButton("Save Image", self)
        self.save_button.clicked.connect(self.save_image)

        # Create a QScrollArea to allow scrolling
        scroll_area = QScrollArea()
        scroll_area.setWidget(self.image_label)
        scroll_area.setWidgetResizable(True)
        scroll_area.setMaximumSize(1000, 1200)

        layout = QVBoxLayout()
        layout.addWidget(scroll_area, 0, Qt.AlignHCenter)

        button_layout = QHBoxLayout()
        button_layout.addStretch()
        button_layout.addWidget(self.save_button)
        button_layout.addStretch()
        layout.addLayout(button_layout)

        self.setLayout(layout)
        self.adjustSize()

    def open_save_image_dialog(self, event):
        options = QFileDialog.Options()
        options |= QFileDialog.ReadOnly
        file_name, _ = QFileDialog.getOpenFileName(self, "Open Image", "", "Images (*.png *.xpm *.jpg *.bmp *.gif *.svg);;All Files (*)", options=options)
        if file_name:
            pixmap = QPixmap(file_name)
            self.text_editor.insert_image(file_name)
            image_window = ImageWindow(pixmap, self)
            # Create the Save button and connect its signal
            image_window.save_button.clicked.connect(image_window.save_image)
            # Add the Save button to the layout
            image_window.layout().addWidget(image_window.save_button)
            image_window.exec_()

    def save_image(self):
        """
        Open a save dialog and save the image to the selected location.
        """
        file_name, _ = QFileDialog.getSaveFileName(
            self, "Save Image", "", "Images (*.png *.xpm *.jpg)"
        )
        if file_name:
            self.original_pixmap.save(file_name)

              
if __name__ == "__main__":
    app = QApplication(sys.argv)
    editor = WordEditor()
    editor.text_editor.textChangedSignal.connect(editor.update_word_count)
    editor.show()
    sys.exit(app.exec_())

